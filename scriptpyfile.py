# -*- coding: utf-8 -*-
import docx
import docxpy
import string
import re
import os
from time import sleep
import csv



def book_info(file_path, name, edition,subheading,isbn,authors,MOTS,REF,Dat,Der):
    fieldnames = ['TITRE DE LA SOLUTION', 'SOMMAIRE', 'LIENS UTILES', 'QUESTIONS À POSER AU CITOYEN','DESCRIPTION','MOTS-CLÉS','RÉFÉRENCES','Date de création','Dernière mise à jour']

    with open(file_path, "a" , newline='',encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        # writer.writeheader()
        writer.writerow({
            "TITRE DE LA SOLUTION": name,
            "SOMMAIRE": edition,
            "LIENS UTILES": subheading,
            "QUESTIONS À POSER AU CITOYEN":isbn,
            "DESCRIPTION":authors,
            "MOTS-CLÉS":MOTS,  
            "RÉFÉRENCES":REF,
            "Date de création":Dat,
            "Dernière mise à jour":Der       
})
j = 0
k = 1
j += k
a = []
b = []
hypertext = ""
left = 0

for i in range(0,228):
    try:
        zee =r'C:\Users\ahmad\Desktop\Projects\wordtocsv\New folder\1 ({}).docx'.format(j)
        print (zee)
        docpx = docxpy.DOCReader(zee)
        docpx.process()  # process file
        hyperlinks = docpx.data['links']
        for z in hyperlinks:
            z1 = z[1]
            z2 = z[0].decode('utf-8')
            hypertext += " {} {} , ".format(z2,z1)
        hyper = hypertext
        doc = docx.Document(zee)
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    d = cell.text
                    a.append(d)
#                     print (cell.text)
                    for paragraph in cell.paragraphs:
                        pass
        book_info("output.csv", a[1],a[3],hyper,a[7],a[9],a[11],a[13],a[16],a[17])
#         print (hyper)
        hyper = ""
        j +=k
        a = []
        hypertext = ""
    except:
        pass

    
    
    